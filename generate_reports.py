from operator import attrgetter

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Side, Border, Font, Alignment
from openpyxl.styles.borders import BORDER_THIN
from openpyxl.worksheet.datavalidation import DataValidation
from setuptools.namespaces import flatten

from boomer_utils import serialize_yes_or_no, adjust_cell_sizes
from models import Participant, Registration, School

MASTER_REPORT = "output/Master.Report.xlsx"
JUDGE_REPORT = "output/Judge.Report.docx"
PARTICPANTS_SHEETS = "output/participants_by_school"
REGULAR_FEE = 12
LATE_FEE = 15


def generate_master_report(events, schools):
    print("GENERATING MASTER REPORT")
    print(F"${REGULAR_FEE} REGULAR / ${LATE_FEE} LATE")

    workbook = Workbook()
    event_worksheet = workbook.active
    event_worksheet.title = 'Events'
    for event in events:
        event_worksheet.append([event.name, len(event.registrations)])
    adjust_cell_sizes(event_worksheet)

    school_worksheet = workbook.create_sheet('Schools')
    school_worksheet.append([
        "School",
        "Regular Registrations",
        "Late Registrations",
        "Registration Fees",
        "Total Enrolled",
        "Rookie Teacher",
        "Rookie School",
        "Attending State"
    ])
    for school in schools:
        school_worksheet.append([
            school.name,
            school.regular_registrations,
            school.late_registrations,
            school.regular_registrations * REGULAR_FEE + school.late_registrations * LATE_FEE,
            school.total_enrolled,
            serialize_yes_or_no(school.rookie_teacher),
            serialize_yes_or_no(school.rookie_school),
            serialize_yes_or_no(school.attending_state)
        ])
    adjust_cell_sizes(school_worksheet)

    workbook.save(MASTER_REPORT)


def generate_judge_report(events):
    events_report = Document()
    sections = events_report.sections 
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.left_margin = Inches(0.5)

    for event in events:
        print("Writing", event.name)
        events_report.add_heading(event.name, 0)
        table = events_report.add_table(rows=1, cols=2)
        if event.max_groups > 0:
            create_group_table(event, table)
        else:
            create_individual_table(event, table)
        table.style = 'Table Grid'
        events_report.add_page_break()

    events_report.save(JUDGE_REPORT)
    print("Saved", JUDGE_REPORT)


def create_group_table(event, table):
    header_cells = table.rows[0].cells
    header_cells[0].text = "School"
    header_cells[1].text = "Participants"

    for cell in header_cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True

    registrations = event.registrations.order_by(Registration.school)
    for registration in registrations:
        row_cells = table.add_row().cells
        school = list(registration.participants)[0].school
        row_cells[0].text = school.name
        participants = registration.participants.order_by(Participant.name)
        row_cells[1].text = '\n'.join(p.name for p in participants)


def create_individual_table(event, table):
    header_cells = table.rows[0].cells
    header_cells[0].text = "Participant"
    header_cells[1].text = "School"

    for cell in header_cells:
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.bold = True

    participants = sorted(event.registrations.participants, key=attrgetter('name'))
    for participant in participants:
        row_cells = table.add_row().cells
        row_cells[0].text = participant.name
        row_cells[1].text = participant.school.name

def generate_participants_sheet(event):
    if len(event.participants) > 0:
        print("Generating", event.name, "participants by school sheet")
    else:
        print("No registrations for", event.name)
        return
    workbook = Workbook()
    worksheet = workbook.active
    
    participants_list = []
    
    for participant in event.participants:
        participants_list.append(participant.name)
        
    participants_list.sort()
    for participant in participants_list:
        worksheet.append([participant,])
    
    participants_sheet = F"{PARTICPANTS_SHEETS}/School.{event.name}.xlsx"
    workbook.save(participants_sheet)
